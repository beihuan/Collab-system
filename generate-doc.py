#!/usr/bin/env python3
"""
Generate the comprehensive collaboration system design document as DOCX.
"""

import sys
sys.path.insert(0, '/usr/lib/python3/dist-packages')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ===== Page Setup =====
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ===== Style Setup =====
style = doc.styles['Normal']
font = style.font
font.name = 'SimHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_font = heading_style.font
    heading_font.name = 'SimHei'
    heading_font.bold = True
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    if level == 1:
        heading_font.size = Pt(22)
        heading_font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    elif level == 2:
        heading_font.size = Pt(16)
        heading_font.color.rgb = RGBColor(0x2D, 0x3A, 0x4A)
    elif level == 3:
        heading_font.size = Pt(13)
        heading_font.color.rgb = RGBColor(0x3D, 0x5A, 0x80)

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'SimHei'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        # Header background
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '2D3A4A')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'SimHei'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            # Alternating row colors
            if row_idx % 2 == 0:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F5F7FA')
                shading.set(qn('w:val'), 'clear')
                cell._tc.get_or_add_tcPr().append(shading)
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    
    doc.add_paragraph()
    return table

def add_code_block(doc, code, language=''):
    """Add a code block with formatting."""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x3A, 0x4A)
    # Add background shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F0F2F5')
    shading.set(qn('w:val'), 'clear')
    p._p.get_or_add_pPr().append(shading)

# ===== COVER PAGE =====
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AI Agent 协作项目管理系统')
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
run.font.name = 'SimHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('完整设计方案 v1.0')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x3D, 0x5A, 0x80)
run.font.name = 'SimHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_paragraph()
doc.add_paragraph()

info_items = [
    ('日期', '2026-04-20'),
    ('版本', '1.0'),
    ('团队规模', '3人'),
    ('项目数量', '2个MVP项目'),
    ('技术栈', 'OpenClaw + 飞书CLI + 飞书多维表格'),
]

for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}: ')
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run = p.add_run(value)
    run.font.size = Pt(12)
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_page_break()

# ===== TABLE OF CONTENTS =====
doc.add_heading('目录', level=1)

toc_items = [
    '1. 系统概述',
    '2. 架构设计',
    '3. 飞书多维表格设计',
    '4. PM Agent 设计',
    '5. 开发者 Agent 设计',
    '6. 协作流程设计',
    '7. 人类确认机制',
    '8. 苏格拉底式追问机制',
    '9. 部署指南',
    '10. 迭代路线图',
]

for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_page_break()

# ===== CHAPTER 1: System Overview =====
doc.add_heading('1. 系统概述', level=1)

doc.add_heading('1.1 问题定义', level=2)

p = doc.add_paragraph()
p.add_run('三人开发团队推进两个MVP项目，当前面临的核心痛点可以拆解为三个层次：').font.name = 'SimHei'

problems = [
    ('信息孤岛', '三人各自用Agent工作，产出分散在各自的Agent会话中，没有统一的信息汇聚点。每天线上会议时，每个人需要从自己的Agent会话中提取关键信息，但由于表达能力不足，关键信息容易遗漏或失真。这种信息不对称导致会议效率低下，经常出现"我以为你已经做完了"或"这个任务不是我来做吗"的误解。'),
    ('协调缺失', '缺乏一个"全局视角"的角色来统筹两个项目的任务分配、优先级排序和依赖关系管理。每个人可能在自己认为重要的事情上投入，但缺乏跨项目、跨人的优先级对齐。特别是当两个项目出现资源冲突时，没有明确的决策机制来决定哪个项目优先。'),
    ('流程缺失', '没有标准化的进度上报和任务接收流程。即使每天开会，也缺少结构化的"输入-处理-输出"闭环，导致会议效率低、对齐不彻底。进度信息散落在各处，无法形成可追溯的项目历史。'),
]

for title, desc in problems:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}：')
    run.bold = True
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run = p.add_run(desc)
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_heading('1.2 系统目标', level=2)

goals = [
    ('Single Source of Truth', '飞书多维表格作为项目进度的唯一真实来源，所有任务状态、优先级、依赖关系都在表格中集中管理，消除信息分散的问题。'),
    ('自动化协调', 'PM Agent作为产品经理兼项目经理，自动规划任务、分配工作、跟踪进度，减少人工协调的负担。'),
    ('结构化沟通', '通过飞书群组IM和飞书文档实现结构化的进度上报和任务指派，确保信息完整、格式统一、可追溯。'),
    ('人类在环', '所有关键操作需人类确认，PM Agent通过苏格拉底式追问持续学习项目上下文，确保Agent的决策始终符合人类意图。'),
    ('深度协作', '开发者Agent自动关联代码提交、PR状态，主动上报阻塞，实现从"被动报告"到"主动协作"的转变。'),
]

for i, (title, desc) in enumerate(goals, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {title}：')
    run.bold = True
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run = p.add_run(desc)
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_heading('1.3 团队角色', level=2)

add_table(doc, 
    ['角色', '人员', '职责', 'Agent运行环境'],
    [
        ['产品经理 + AI Coding开发', 'Person A（你）', '统筹两个项目、全栈开发、项目管理', 'OpenClaw（你的电脑）'],
        ['AI Coding开发', 'Person B', '两个项目的全栈开发', 'OpenClaw（Person B的电脑）'],
        ['后端开发', 'Person C', '后端开发、数据库维护', 'OpenClaw（Person C的电脑）'],
    ]
)

doc.add_heading('1.4 技术栈', level=2)

add_table(doc,
    ['组件', '技术', '用途'],
    [
        ['项目管理', '飞书多维表格 (Bitable)', '任务进度的Single Source of Truth'],
        ['文档协作', '飞书云文档', '任务指派、进度报告、会议纪要'],
        ['即时通信', '飞书群组IM', 'Agent间实时通信、通知'],
        ['CLI工具', '飞书CLI (lark-cli)', '操作多维表格、文档、IM'],
        ['Agent框架', 'OpenClaw', '运行PM Agent和开发者Agent'],
        ['定时调度', 'OpenClaw Cron Job', '定时触发工作流'],
        ['周期检查', 'OpenClaw Heartbeat', '周期性状态检查'],
        ['自定义工作流', 'OpenClaw Skill', '封装各角色的专业能力'],
    ]
)

# ===== CHAPTER 2: Architecture =====
doc.add_heading('2. 架构设计', level=1)

doc.add_heading('2.1 系统架构图', level=2)

p = doc.add_paragraph()
p.add_run('整个系统采用三层架构，飞书多维表格作为数据层，PM Agent作为协调层，开发者Agent作为执行层：').font.name = 'SimHei'

add_code_block(doc, """
┌─────────────────────────────────────────────────────────────────┐
│                    飞书多维表格 (Single Source of Truth)          │
│          项目任务进度、状态、优先级、依赖关系、验收标准            │
└────────────────────────────┬────────────────────────────────────┘
                             │ 飞书CLI 读写
                             ▼
┌─────────────────────────────────────────────────────────────────┐
│                    PM Agent (OpenClaw - Person A的电脑)           │
│  ┌──────────────┐ ┌──────────────┐ ┌──────────────────────┐    │
│  │ 任务规划      │ │ 进度审核      │ │ 会议辅助              │    │
│  │ morning-planner│ │ progress-    │ │ meeting-assistant    │    │
│  │              │ │ reviewer     │ │                      │    │
│  └──────────────┘ └──────────────┘ └──────────────────────┘    │
│  ┌──────────────┐ ┌──────────────┐ ┌──────────────────────┐    │
│  │ 表格管理      │ │ 上下文管理    │ │ 苏格拉底学习          │    │
│  │ bitable-mgr  │ │ context-mgr  │ │ socratic-learner     │    │
│  └──────────────┘ └──────────────┘ └──────────────────────┘    │
│                                                                 │
│  Cron: 08:00规划 / 09:30会议 / 19:00审核                        │
│  Heartbeat: 每30分钟检查阻塞/逾期/报告                           │
└──────────┬─────────────────────────────────┬───────────────────┘
           │ 飞书群组IM + 飞书文档             │
           ▼                                  ▼
┌──────────────────────┐  ┌──────────────────────────────────────┐
│ Person B Agent        │  │ Person C Agent                        │
│ (OpenClaw - B的电脑)  │  │ (OpenClaw - C的电脑)                  │
│ ┌──────────────────┐ │  │ ┌──────────────────┐                  │
│ │ 进度上报          │ │  │ │ 进度上报          │                  │
│ │ progress-reporter │ │  │ │ progress-reporter │                  │
│ ├──────────────────┤ │  │ ├──────────────────┤                  │
│ │ 任务接收          │ │  │ │ 任务接收          │                  │
│ │ task-receiver     │ │  │ │ task-receiver     │                  │
│ ├──────────────────┤ │  │ ├──────────────────┤                  │
│ │ 阻塞通知          │ │  │ │ 阻塞通知          │                  │
│ │ blocker-notifier  │ │  │ │ blocker-notifier  │                  │
│ ├──────────────────┤ │  │ ├──────────────────┤                  │
│ │ 代码活动追踪      │ │  │ │ 代码活动追踪      │                  │
│ │ code-activity     │ │  │ │ code-activity     │                  │
│ └──────────────────┘ │  │ └──────────────────┘                  │
│ Cron: 09:00接收/18:00上报│ │ Cron: 09:00接收/18:00上报            │
│ Heartbeat: 每30分钟    │  │ Heartbeat: 每30分钟                  │
└──────────────────────┘  └──────────────────────────────────────┘
""")

doc.add_heading('2.2 通信架构', level=2)

p = doc.add_paragraph()
p.add_run('所有Agent间的通信通过飞书群组IM进行，飞书文档作为结构化信息的载体。这种设计有以下优势：').font.name = 'SimHei'

benefits = [
    '统一通信渠道：所有Agent和人类都在同一个飞书群组中，信息透明可追溯',
    '结构化文档：任务指派和进度报告使用飞书文档，格式统一、内容完整',
    '实时通知：阻塞报告和紧急事项通过IM消息即时传达',
    '人类可介入：人类可以随时在群组中查看Agent的对话，也可以直接参与讨论',
    '异步兼容：即使某人的Agent离线，消息也会保留在群组中，上线后可以读取',
]

for b in benefits:
    p = doc.add_paragraph(b, style='List Bullet')
    for run in p.runs:
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_heading('2.3 数据流', level=2)

add_code_block(doc, """
每日数据流：

08:00  PM Agent Cron触发
       │
       ├─→ 读取多维表格当前状态
       ├─→ 分析任务优先级和依赖
       ├─→ 生成任务指派文档（每人一份）
       ├─→ 人类确认 → 发送到飞书群组
       │
09:00  开发者Agent Cron触发
       │
       ├─→ 读取飞书群组中的任务指派
       ├─→ 展示给人类确认
       ├─→ 发送确认回执
       │
09:30  PM Agent Cron触发 → 生成会议议程
       │
10:00  每日对齐会议（人类主导）
       │
...工作时间...  开发者Agent Heartbeat监控
       │
       ├─→ 每30分钟检查代码活动
       ├─→ 检测阻塞 → 人类确认 → 通知PM Agent
       │
18:00  开发者Agent Cron触发
       │
       ├─→ 收集今日代码活动
       ├─→ 生成进度报告 → 人类确认 → 发送到飞书群组
       │
19:00  PM Agent Cron触发
       │
       ├─→ 读取所有进度报告
       ├─→ 分析进度偏差
       ├─→ 提出多维表格更新 → 人类确认 → 执行更新
       ├─→ 如有不确定 → 苏格拉底式追问
       │
会后   人类提供会议纪要
       │
       └─→ PM Agent处理纪要 → 更新任务和上下文
""")

# ===== CHAPTER 3: Bitable Design =====
doc.add_heading('3. 飞书多维表格设计', level=1)

doc.add_heading('3.1 表格结构', level=2)

p = doc.add_paragraph()
p.add_run('采用一个多维表格应用、一个数据表、用"项目"字段区分两个项目的方案。这种设计便于跨项目查看和资源协调，同时保持数据结构简单。').font.name = 'SimHei'

add_table(doc,
    ['字段名', '字段类型', '可选值/格式', '必填', '说明'],
    [
        ['task_id', '自动编号', 'TASK-{4位数字}', '自动', '唯一任务标识'],
        ['project', '单选', '项目A, 项目B', '是', '所属项目'],
        ['title', '文本', '自由文本', '是', '任务标题'],
        ['description', '文本', '自由文本', '是', '详细任务描述'],
        ['assignee', '单选', 'Person A, Person B, Person C', '是', '负责人'],
        ['status', '单选', '待开始, 进行中, 待评审, 已完成, 已阻塞', '是', '任务状态'],
        ['priority', '单选', 'P0-紧急, P1-高, P2-中, P3-低', '是', '优先级'],
        ['dependencies', '文本', '逗号分隔的TASK-ID', '否', '前置依赖任务'],
        ['due_date', '日期', 'YYYY-MM-DD', '是', '截止日期'],
        ['acceptance_criteria', '文本', '自由文本', '是', '验收标准'],
        ['milestone', '单选', 'MVP-核心功能, MVP-完善功能, MVP-上线准备', '是', '所属里程碑'],
        ['pr_link', '超链接', 'URL', '否', '关联PR链接'],
        ['blocker', '文本', '自由文本', '否', '阻塞描述'],
        ['estimated_hours', '数字', '正整数', '否', '预估工时(小时)'],
        ['actual_hours', '数字', '正整数', '否', '实际工时(小时)'],
        ['progress', '单选', '0%, 25%, 50%, 75%, 90%, 100%', '是', '进度百分比'],
        ['risk_level', '单选', '无风险, 低风险, 中风险, 高风险', '否', '风险等级'],
    ]
)

doc.add_heading('3.2 状态流转规则', level=2)

add_code_block(doc, """
待开始 ──→ 进行中 ──→ 待评审 ──→ 已完成
  │          │
  │          ↓
  └────→ 已阻塞 ──→ 进行中 (阻塞解除后)
""")

p = doc.add_paragraph()
p.add_run('状态流转说明：').bold = True
states = [
    ('待开始 → 进行中', '开发者开始工作时，由PM Agent或开发者Agent更新'),
    ('进行中 → 待评审', '开发者完成开发并提交PR后更新'),
    ('待评审 → 已完成', 'PR合并且验收标准全部满足后更新'),
    ('进行中 → 已阻塞', '遇到无法自行解决的问题时更新，需填写blocker字段'),
    ('已阻塞 → 进行中', '阻塞解除后更新，需清空blocker字段'),
]
for transition, desc in states:
    p = doc.add_paragraph()
    run = p.add_run(f'{transition}：')
    run.bold = True
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run = p.add_run(desc)
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_heading('3.3 优先级定义', level=2)

add_table(doc,
    ['优先级', '定义', '响应时间', '示例'],
    [
        ['P0-紧急', '影响上线或阻塞他人的关键任务', '立即处理', '线上bug、阻塞他人的接口'],
        ['P1-高', '里程碑核心功能', '当日处理', 'MVP必须功能、关键API'],
        ['P2-中', '重要但非紧急', '本周处理', '功能优化、非核心功能'],
        ['P3-低', '优化和改进类', '有空处理', '代码重构、文档完善'],
    ]
)

# ===== CHAPTER 4: PM Agent Design =====
doc.add_heading('4. PM Agent 设计', level=1)

doc.add_heading('4.1 Skill 清单', level=2)

add_table(doc,
    ['Skill名称', '功能', '触发方式', '核心能力'],
    [
        ['pm-bitable-manager', '多维表格CRUD操作', '被其他Skill调用', '读写记录、分层确认、批量操作'],
        ['pm-morning-planner', '每日早晨任务规划', 'Cron 08:00', '任务分配、优先级排序、依赖检查'],
        ['pm-progress-reviewer', '每日晚间进度审核', 'Cron 19:00', '偏差分析、风险识别、表格更新'],
        ['pm-meeting-assistant', '会议辅助', 'Cron 09:30 + 手动', '议程生成、纪要处理、行动项跟踪'],
        ['pm-context-manager', '项目上下文管理', '初始化 + 手动', '代码库阅读、文档维护、理解更新'],
        ['pm-socratic-learner', '苏格拉底式学习', '纠正时 + 信息缺失时', '深层追问、理解更新、上下文补充'],
    ]
)

doc.add_heading('4.2 pm-morning-planner 详解', level=2)

p = doc.add_paragraph()
p.add_run('这是PM Agent最核心的Skill，负责每天早晨为每个团队成员生成个性化的任务指派文档。它的工作流程如下：').font.name = 'SimHei'

steps = [
    ('读取多维表格', '获取所有任务的当前状态，包括进行中、待开始、已阻塞的任务，以及每个人的当前工作负载。'),
    ('分析任务优先级', '根据P0-P3优先级排序，同时考虑依赖关系——如果高优先级任务依赖其他任务，被依赖的任务也需要提升优先级。'),
    ('检查依赖关系', '确保不分配依赖未满足的任务。如果依赖正在进行中，标注预计完成时间；如果依赖被阻塞，标注阻塞原因。'),
    ('负载均衡', '检查每个人的当前任务数量和预估工时，避免单人日工作量超过6小时。如果某人过载，考虑将P2/P3任务重新分配。'),
    ('生成指派文档', '为每个人生成包含详细任务描述、验收标准、依赖关系、优先级排序的指派文档。'),
    ('人类确认', '展示指派方案给人类确认，允许修改后发送到飞书群组。'),
]

for i, (step, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {i}: {step}')
    run.bold = True
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    p = doc.add_paragraph(desc)
    for run in p.runs:
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_heading('4.3 任务指派文档示例', level=2)

add_code_block(doc, """
📬 任务指派 — Person B — 2026-04-20

━━━━━━━━━━━━━━━━━━━━━━━━━━━━

🔴 P0 — 紧急

[TASK-0001] 实现用户登录API
├─ 项目: 项目A
├─ 描述: 实现基于JWT的用户登录API，支持邮箱/手机号登录，
│       集成第三方OAuth（微信、Google）。此接口是支付模块
│       的前置依赖，必须在今天完成。
├─ 验收标准:
│  1. 邮箱登录接口返回有效JWT token
│  2. 手机号+验证码登录正常工作
│  3. 微信OAuth回调正确处理
│  4. 单元测试覆盖率>80%
│  5. API文档已更新
├─ 依赖: TASK-0005 (数据库用户表已创建) ✅ 已完成
├─ 截止日期: 2026-04-20 (今天!)
└─ 备注: 后端同学已确认测试环境数据库已就绪

━━━━━━━━━━━━━━━━━━━━━━━━━━━━

🟠 P1 — 高

[TASK-0023] 设计权限管理模块
├─ 项目: 项目A
├─ 描述: 基于RBAC模型设计权限管理模块的数据模型和API接口
├─ 验收标准:
│  1. 数据模型设计文档已评审通过
│  2. API接口列表已定义
│  3. 与现有用户体系兼容
├─ 依赖: TASK-0001 (用户登录API) 🔄 进行中
├─ 截止日期: 2026-04-25
└─ 备注: 可在TASK-0001完成后开始

━━━━━━━━━━━━━━━━━━━━━━━━━━━━

📌 待关注事项
- TASK-0007 (第三方API文档) 仍在等待中，如今天收到请优先处理
""")

doc.add_heading('4.4 pm-progress-reviewer 详解', level=2)

p = doc.add_paragraph()
p.add_run('晚间进度审核Skill负责收集和分析当天的进度报告，识别偏差和风险，并提出多维表格更新建议。它的核心能力在于偏差检测：').font.name = 'SimHei'

add_table(doc,
    ['偏差类型', '检测逻辑', '严重程度', '处理方式'],
    [
        ['进度落后', '报告进度 < 预期进度20%以上', '中', '标记风险，询问原因'],
        ['进度严重落后', '报告进度 < 预期进度40%以上', '高', '立即通知人类，建议调整计划'],
        ['计划外工作', '提到但不在多维表格中的任务', '低', '建议添加到表格'],
        ['缺失报告', '19:30前未提交报告', '高', '发送提醒，标记风险'],
        ['新阻塞', '报告了表格中没有的阻塞', '中', '更新表格，分析影响'],
        ['截止日期变更', '报告的ETA与表格不一致', '中', '确认后更新表格'],
        ['任务未提及', '表格中"进行中"但报告中未提到', '中', '询问任务状态'],
    ]
)

doc.add_heading('4.5 pm-socratic-learner 详解', level=2)

p = doc.add_paragraph()
p.add_run('苏格拉底式学习是PM Agent持续进步的核心机制。它分为三个追问层次：').font.name = 'SimHei'

add_table(doc,
    ['层次', '类型', '典型问题', '触发场景'],
    [
        ['Level 1', '澄清性', '"你说的XX具体是指什么？"', '术语不清、需求模糊'],
        ['Level 2', '推理性', '"为什么是这样？如果改成YY会怎样？"', '理解决策背后的原因'],
        ['Level 3', '深层', '"这对整体里程碑有什么影响？我们是否需要重新评估？"', '理解全局影响和隐含假设'],
    ]
)

p = doc.add_paragraph()
p.add_run('追问触发场景与示例：').bold = True

socratic_examples = [
    ('人类纠正优先级', 
     '人类: "TASK-0003的优先级应该是P0，不是P1"',
     'PM Agent: "我理解了，TASK-0003应该是P0。让我深入了解一下——为什么这个任务的优先级需要提升？它是否阻塞了其他关键任务？如果提升为P0，对今天其他P0任务的排期有什么影响？我们是否需要推迟某些P1任务？"'),
    ('信息不完整', 
     'PM Agent发现任务描述过于简略',
     'PM Agent: "我注意到TASK-0023的描述比较简略。这个权限管理模块需要支持哪些具体的权限类型？是仅限于功能权限还是也包括数据权限？现有的用户体系中是否已经有角色概念？这些信息会影响数据模型的设计和工时估算。"'),
    ('进度偏差原因不明', 
     '开发者报告进度落后但未说明原因',
     'PM Agent: "TASK-0001的进度从80%降到了60%，能帮我理解一下发生了什么吗？是发现了之前未预料到的问题，还是需求发生了变化？这个问题是否会影响验收标准？如果需要调整计划，我们有哪些选项？"'),
]

for scenario, trigger, response in socratic_examples:
    p = doc.add_paragraph()
    run = p.add_run(f'场景：{scenario}')
    run.bold = True
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    p = doc.add_paragraph(f'触发：{trigger}')
    p = doc.add_paragraph(f'回应：{response}')

# ===== CHAPTER 5: Developer Agent Design =====
doc.add_heading('5. 开发者 Agent 设计', level=1)

doc.add_heading('5.1 Skill 清单', level=2)

add_table(doc,
    ['Skill名称', '功能', '触发方式', '核心能力'],
    [
        ['dev-progress-reporter', '进度上报', 'Cron 18:00', '收集代码活动、生成结构化报告、人类确认'],
        ['dev-task-receiver', '任务接收', 'Cron 09:00', '读取指派文档、展示任务、确认回执'],
        ['dev-blocker-notifier', '阻塞通知', '手动 + Heartbeat检测', '阻塞检测、影响分析、即时通知'],
        ['dev-code-activity-tracker', '代码活动追踪', 'Heartbeat + 被调用', 'Git提交关联、PR状态、代码审查'],
    ]
)

doc.add_heading('5.2 dev-progress-reporter 详解', level=2)

p = doc.add_paragraph()
p.add_run('进度上报是开发者Agent最核心的职责。它通过三个信息源自动收集今日工作情况，然后由人类补充和确认：').font.name = 'SimHei'

sources = [
    ('代码活动', '自动收集今日的git提交、PR创建/更新、代码审查活动，并尝试通过分支名和提交信息关联到多维表格中的任务ID。'),
    ('任务指派', '读取早晨收到的任务指派文档，对比计划与实际完成情况。'),
    ('人工输入', '向开发者展示自动收集的信息，让开发者补充非代码工作（如文档编写、会议参与、技术调研等），并确认最终报告内容。'),
]

for title, desc in sources:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}：')
    run.bold = True
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run = p.add_run(desc)
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_heading('5.3 进度报告示例', level=2)

add_code_block(doc, """
📊 进度汇报 — Person B — 2026-04-20

━━━━━━━━━━━━━━━━━━━━━━━━━━━━

## 当前任务

- [TASK-0001] 实现用户登录API — 进度: 80%
  - 状态: 进行中
  - 今日进展: 完成了邮箱登录和JWT token生成逻辑，
    微信OAuth回调处理进行中
  - 关联提交: a3f2c1d, b7e4a9f
  - 关联PR: #42 (进行中)

## 本周期已完成事项

- [TASK-0015] 修复支付回调bug — ✅ 已完成
  - 产出: PR #40 已合并
  - 实际工时: 3小时 (预估: 2小时)

## 下周期计划事项

- [TASK-0001] 完成微信OAuth回调 + 单元测试
  - 计划进度: 100%
- [TASK-0023] 开始权限管理模块设计
  - 计划进度: 20%

## 风险/卡点

- 🚫 [TASK-0001] 第三方OAuth服务商的沙箱环境不稳定
  - 原因: 沙箱环境频繁返回500错误
  - 影响: 可能影响联调进度，但不影响本地开发
  - 需要协助: 如明天仍不稳定，可能需要联系服务商

## 预计完成时间变更

- [TASK-0001] 原计划 2026-04-20 → 预计 2026-04-22
  - 原因: OAuth沙箱环境不稳定导致联调延迟

## 代码活动

- 提交: 4个commits (a3f2c1d, b7e4a9f, c1d3e5f, d2f4a6b)
- PR: #42 进行中, #40 已合并
- 代码审查: 审查了Person C的PR #41
""")

doc.add_heading('5.4 dev-blocker-notifier 详解', level=2)

p = doc.add_paragraph()
p.add_run('阻塞通知机制支持两种触发方式：手动触发（开发者主动告知Agent遇到阻塞）和自动检测（Heartbeat发现某任务长时间无代码活动）。阻塞报告不仅描述问题，还包含影响分析和临时方案建议，帮助PM Agent快速做出决策。').font.name = 'SimHei'

doc.add_heading('5.5 dev-code-activity-tracker 详解', level=2)

p = doc.add_paragraph()
p.add_run('代码活动追踪是深度协作的关键。它通过分析git提交信息中的任务ID引用（如"feat(TASK-0001): add login API"）、分支命名规则（如"feature/TASK-0001-login"）和PR描述中的任务关联，自动将代码活动与多维表格中的任务关联起来。这种关联使得进度报告不仅仅是主观描述，还有客观的代码活动数据支撑。').font.name = 'SimHei'

p = doc.add_paragraph()
p.add_run('约定：为了使代码活动追踪更准确，团队需要遵循以下约定：').font.name = 'SimHei'

conventions = [
    '提交信息格式：type(TASK-XXXX): description，如 feat(TASK-0001): add JWT login',
    '分支命名：feature/TASK-XXXX-description 或 fix/TASK-XXXX-description',
    'PR描述中引用任务ID：Refs: TASK-XXXX',
    'PR标题包含任务ID：[TASK-XXXX] Feature description',
]

for c in conventions:
    p = doc.add_paragraph(c, style='List Bullet')
    for run in p.runs:
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

# ===== CHAPTER 6: Collaboration Flow =====
doc.add_heading('6. 协作流程设计', level=1)

doc.add_heading('6.1 每日完整工作流', level=2)

p = doc.add_paragraph()
p.add_run('以下是一个完整工作日的协作流程时间线，展示了所有Agent的触发时机和交互过程：').font.name = 'SimHei'

timeline = [
    ('08:00', 'PM Agent', '早晨规划', 'Cron触发pm-morning-planner，读取多维表格，为每人生成任务指派文档，人类确认后发送到飞书群组'),
    ('09:00', '开发者Agent', '任务接收', 'Cron触发dev-task-receiver，读取飞书群组中的任务指派，展示给人类确认，发送确认回执'),
    ('09:30', 'PM Agent', '会议准备', 'Cron触发pm-meeting-assistant，生成每日对齐会议的进度摘要和议程'),
    ('10:00', '人类', '每日会议', '人类根据议程进行讨论，PM Agent不参与会议本身'),
    ('工作中', '开发者Agent', 'Heartbeat监控', '每30分钟检查代码活动和任务进度，检测阻塞'),
    ('随时', '开发者Agent', '阻塞通知', '如遇阻塞，人类确认后发送阻塞报告到飞书群组，PM Agent的Heartbeat会检测到'),
    ('18:00', '开发者Agent', '进度上报', 'Cron触发dev-progress-reporter，收集代码活动和人工输入，生成进度报告，人类确认后发送'),
    ('19:00', 'PM Agent', '晚间审核', 'Cron触发pm-progress-reviewer，收集所有进度报告，分析偏差，提出表格更新，人类确认后执行'),
    ('会后', '人类→PM Agent', '会议纪要', '人类将会议纪要提供给PM Agent，触发会后处理流程'),
]

add_table(doc,
    ['时间', '执行者', '动作', '详细说明'],
    timeline
)

doc.add_heading('6.2 异常流程', level=2)

doc.add_heading('6.2.1 阻塞处理流程', level=3)

add_code_block(doc, """
开发者遇到阻塞
    │
    ├─→ 手动: 告诉Agent "我被阻塞了"
    │    或
    ├─→ 自动: Heartbeat检测到4小时无代码活动
    │
    ↓
开发者Agent收集阻塞详情
    │
    ↓
人类确认阻塞报告
    │
    ↓
发送阻塞报告到飞书群组
    │
    ↓
PM Agent Heartbeat检测到阻塞报告
    │
    ├─→ 分析影响范围（哪些依赖任务受影响）
    ├─→ 评估是否需要调整其他人的任务
    ├─→ 通知人类并建议应对方案
    │
    ↓
人类决策（调整任务/协调资源/联系外部）
    │
    ↓
PM Agent更新多维表格和任务指派
""")

doc.add_heading('6.2.2 逾期任务处理流程', level=3)

add_code_block(doc, """
PM Agent Heartbeat检测到逾期任务
    │
    ↓
通知人类（包含逾期原因分析）
    │
    ├─→ 延期合理: 调整截止日期
    ├─→ 需要支援: 重新分配资源
    ├─→ 需要删减: 调整验收标准或降级优先级
    │
    ↓
PM Agent执行调整（人类确认后）
""")

# ===== CHAPTER 7: Human Confirmation =====
doc.add_heading('7. 人类确认机制', level=1)

doc.add_heading('7.1 三层确认体系', level=2)

p = doc.add_paragraph()
p.add_run('为了平衡自动化效率和人类控制权，系统采用三层确认体系，根据操作的风险等级决定确认级别：').font.name = 'SimHei'

add_table(doc,
    ['确认级别', '操作类型', '确认方式', '示例'],
    [
        ['L0 - 自动执行', '只读操作', '无需确认', '读取多维表格、读取飞书消息、生成草稿'],
        ['L2 - 单次确认', '常规写操作', '终端y/n确认', '更新任务状态、发送任务指派、发送进度报告'],
        ['L3 - 双重确认', '高风险操作', '终端确认+原因说明', '删除任务、修改截止日期、标记任务阻塞、调整优先级'],
    ]
)

doc.add_heading('7.2 批量确认', level=2)

p = doc.add_paragraph()
p.add_run('当PM Agent在晚间审核时需要更新多条记录，系统支持批量确认机制，避免逐条确认的繁琐：').font.name = 'SimHei'

add_code_block(doc, """
即将批量更新多维表格（共5条记录）：

状态变更:
  - [TASK-0001] 状态: 进行中 → 待评审
  - [TASK-0015] 状态: 待开始 → 已完成
  - [TASK-0023] 状态: 待开始 → 进行中, 进度: 20%

⚠️ 高风险变更:
  - [TASK-0007] 截止日期: 2026-04-20 → 2026-04-25 (延迟5天)
    原因: 第三方API文档延迟交付

确认以上所有变更？(y/n/edit/split)
- y: 确认全部
- n: 取消全部
- edit: 修改特定项
- split: 分别确认每项
""")

# ===== CHAPTER 8: Socratic Mechanism =====
doc.add_heading('8. 苏格拉底式追问机制', level=1)

doc.add_heading('8.1 触发条件', level=2)

add_table(doc,
    ['触发场景', '触发方式', '追问层次', '示例'],
    [
        ['人类纠正PM Agent的理解', '被动触发', 'Level 2-3', '人类说"优先级不对"，追问为什么和影响'],
        ['人类纠正PM Agent的决策', '被动触发', 'Level 2-3', '人类说"不该分配给B"，追问原因和更好方案'],
        ['任务信息不完整', '主动检测', 'Level 1-2', '任务缺少验收标准，追问具体标准'],
        ['进度报告含糊', '主动检测', 'Level 1-2', '报告说"基本完成"，追问具体完成度'],
        ['依赖关系不明确', '主动检测', 'Level 2', '任务间依赖未标注，追问依赖详情'],
        ['项目背景有盲区', '主动检测', 'Level 2-3', '发现不熟悉的技术或业务概念，追问背景'],
    ]
)

doc.add_heading('8.2 追问策略', level=2)

p = doc.add_paragraph()
p.add_run('追问不是无限制的，需要遵循以下策略来平衡信息获取和人类耐心：').font.name = 'SimHei'

strategies = [
    ('优先级排序', '每次追问最多5个问题，优先问对项目理解影响最大的问题。如果人类显得不耐烦，减少到1-2个最关键的问题。'),
    ('渐进深入', '第一次交互问2-3个Level 2问题，后续追问1-2个Level 3问题。不要一上来就问最深层的问题。'),
    ('推断优先', '如果能从已有信息推断出答案，先推断再确认，而不是直接问。比如人类说"这个功能和XX类似"，PM Agent应该先基于对XX的理解做出推断，然后问"我的理解是...，对吗？"'),
    ('记录避免重复', '所有追问的结论都记录在项目上下文文档中，避免重复问同一个问题。'),
    ('紧急情况让步', '如果决策紧急，先做临时决定，标注不确定性，事后再追问。不要因为追问而延误关键决策。'),
]

for title, desc in strategies:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}：')
    run.bold = True
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run = p.add_run(desc)
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_heading('8.3 项目上下文持久化', level=2)

p = doc.add_paragraph()
p.add_run('所有通过苏格拉底式追问获得的项目理解，都必须持久化到项目上下文文档中。这个文档是PM Agent的"记忆"，确保即使对话上下文被清除，PM Agent也能通过读取这个文档恢复对项目的理解。文档结构包括：项目概述、技术架构、核心模块、关键决策记录、已知约束、当前焦点、跨项目信息等。每次更新都记录日期、原因和变更内容，形成可追溯的理解演进历史。').font.name = 'SimHei'

# ===== CHAPTER 9: Deployment =====
doc.add_heading('9. 部署指南', level=1)

doc.add_heading('9.1 前置条件', level=2)

prereqs = [
    '飞书企业版账号，具有创建多维表格和群组的权限',
    '飞书CLI (lark-cli) 已安装并完成机器人身份认证',
    'OpenClaw 已安装在三台电脑上',
    '三个项目代码仓库可在本地访问',
    '飞书机器人应用已创建，获得App ID和App Secret',
]

for p_text in prereqs:
    p = doc.add_paragraph(p_text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_heading('9.2 部署步骤', level=2)

deploy_steps = [
    ('创建飞书资源', '创建多维表格应用并按第3章的表结构配置字段；创建三人+机器人的飞书群组；记录APP_TOKEN、TABLE_ID、CHAT_ID'),
    ('配置飞书CLI', '在三台电脑上分别安装和认证飞书CLI，使用机器人身份；验证可以读写多维表格和发送群消息'),
    ('部署PM Agent Skills', '将pm-agent/skills/下的6个Skill文件夹复制到OpenClaw的skills目录；编辑各Skill的references/feishu-config.md和references/repo-config.md'),
    ('部署开发者Agent Skills', '将dev-agent/skills/下的4个Skill文件夹复制到各自的OpenClaw skills目录；编辑references/feishu-config.md和references/identity.md'),
    ('配置Cron Jobs', '按4.3节的调度配置在各Agent的openclaw.json中添加Cron任务'),
    ('配置Heartbeat', '按各Agent的HEARTBEAT.md配置心跳检查项'),
    ('初始化PM Agent', '与PM Agent进行初始化对话，让它阅读代码库和文档，通过苏格拉底式追问建立项目上下文'),
    ('端到端测试', '手动触发各Skill，验证完整流程是否正常工作'),
]

for i, (step, desc) in enumerate(deploy_steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {i}: {step}')
    run.bold = True
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    p = doc.add_paragraph(desc)
    for run in p.runs:
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

# ===== CHAPTER 10: Roadmap =====
doc.add_heading('10. 迭代路线图', level=1)

doc.add_heading('10.1 Phase 1: 最小可用系统（第1周）', level=2)

p = doc.add_paragraph()
p.add_run('目标：基本的任务管理和进度上报闭环').bold = True

phase1 = [
    '飞书多维表格创建和字段配置',
    'pm-bitable-manager Skill（基础CRUD）',
    'pm-morning-planner Skill（简化版：仅列出今日任务）',
    'dev-progress-reporter Skill（基础版：手动输入进度）',
    'dev-task-receiver Skill（读取任务指派）',
    '飞书群组基本通信',
    'L2确认机制',
]

for item in phase1:
    p = doc.add_paragraph(f'[ ] {item}')
    for run in p.runs:
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_heading('10.2 Phase 2: 自动化增强（第2-3周）', level=2)

p = doc.add_paragraph()
p.add_run('目标：减少手动操作，增加自动化').bold = True

phase2 = [
    'pm-progress-reviewer Skill（自动分析进度偏差）',
    'pm-meeting-assistant Skill（会前摘要+会后处理）',
    'dev-blocker-notifier Skill（阻塞检测和告警）',
    'dev-code-activity-tracker Skill（git提交关联）',
    'Heartbeat检查项上线',
    'L3审批机制',
]

for item in phase2:
    p = doc.add_paragraph(f'[ ] {item}')
    for run in p.runs:
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_heading('10.3 Phase 3: 智能化增强（第4-6周）', level=2)

p = doc.add_paragraph()
p.add_run('目标：PM Agent越来越懂项目').bold = True

phase3 = [
    'pm-context-manager Skill（项目上下文自动维护）',
    'pm-socratic-learner Skill（深层追问机制）',
    '项目上下文文档持久化和版本管理',
    '基于历史数据的工时估算优化',
    '风险预警（基于进度趋势预测逾期）',
]

for item in phase3:
    p = doc.add_paragraph(f'[ ] {item}')
    for run in p.runs:
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_heading('10.4 Phase 4: 高级特性（后续迭代）', level=2)

phase4 = [
    '跨项目资源冲突检测',
    '自动生成周报/月报',
    '与CI/CD集成（构建状态关联）',
    '飞书审批流集成（替代终端确认）',
    '多维表格视图自动配置',
]

for item in phase4:
    p = doc.add_paragraph(f'[ ] {item}')
    for run in p.runs:
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

# ===== Save =====
output_path = '/home/z/my-project/download/collab-system/AI-Agent协作项目管理系统-设计方案.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
